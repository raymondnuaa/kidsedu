import time,sys,os,re
from docx import Document
from docx.shared import Inches

def PareLinesInfo(resFileName):
    fileQues = open(resFileName, "r")
    oneDayQues = ''
    linesInfo = {}
    oneDay = []
    dayKey = ''

    i = 0
    for line in fileQues:
        if(len(line) > 20):
            oneDayQues = line.decode('utf8')
            oneDay.append(oneDayQues)            
            if(dayKey == ''):
                datimeInfo = re.search('(\d+.\d+)', oneDayQues)
                if(datimeInfo and datimeInfo.group(1)):
                    dayKey = datimeInfo.group(1)
                    
        elif(line.startswith(';')):
            linesInfo[dayKey] = oneDay            
            oneDay = []
            dayKey = ''
            i += 1     
        
        
    return linesInfo

        
def GenFinalDocx(linesInfo, docFileName) :
    document = Document()    
    index = 1
    
    #for key in sorted(linesInfo.iterkeys()):
    for i in range(1, 13):
        monthDoc = Document()
        inMonthCnt = 1
        for j in range(1, 32):
            key = '%d.%d' % (i, j)
            if (key in linesInfo.keys()):
                k = 0
                for infoLine in linesInfo[key]:
                    if(k == 0):
                        indexLine = '%d - (%d/%d): ' % (index, i, j)
                        document.add_heading(indexLine + infoLine, level=2)
                        monthDoc.add_heading(str(inMonthCnt) + ': ' + infoLine, level=2)
                    elif(infoLine.startswith('full')):
                        document.add_picture(infoLine[:-1],  width=Inches(6.0))  
                        monthDoc.add_picture(infoLine[:-1],  width=Inches(6.0))  
                    else:
                        document.add_paragraph(infoLine)
                        monthDoc.add_paragraph(infoLine)
                    k += 1
                  
                inMonthCnt += 1
                index += 1
                
        monthDoc.save('%s-%d.docx' % (docFileName, i))        
    document.save(docFileName + '.docx')    

if __name__ == "__main__":
    linesInfo = PareLinesInfo("itemQues.txt")
    GenFinalDocx(linesInfo, 'AoshuQuestion')
    
    linesInfo = PareLinesInfo("itemAnws.txt")
    GenFinalDocx(linesInfo, 'AoshuAnwser')
